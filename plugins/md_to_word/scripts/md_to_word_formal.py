#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# 依赖安装：
# pip uninstall docx -y 2>nul
# pip install python-docx
"""
md_to_word_formal.py
基于参考 docx 样式（方案/规划文档排版）将 Markdown 文件转换为 Word 文档。

样式规格来源：
  参考文件: 响应人工智能管理框架的工业AI工作规划.docx
  中文字体取自参考 docx 的 w:eastAsia 属性，西文/数字使用 Times New Roman。

格式映射：
  #       → 文档标题   (中:宋体           / 西:Times New Roman  20pt Bold 居中  32pt行距 段后18pt)
  ##      → Heading 1  (中:黑体           / 西:Times New Roman  16pt Bold 左对齐 28pt行距)
  ###     → Heading 2  (中:楷体           / 西:Times New Roman  16pt Bold 左对齐 28pt行距)
  ####    → Heading 3  (中:宋体           / 西:Times New Roman  14pt Bold 左对齐 28pt行距 首行缩进1.42cm)
  #####   → Heading 4  (中:仿宋           / 西:Times New Roman  14pt Bold 左对齐 段前9pt  首行缩进1.42cm)
  正文     → 宋体 14pt 两端对齐 首行缩进1.13cm 28pt行距
  表格     → 三线表（顶线/底线 1.5pt，栏目线 0.75pt） 宋体 12pt 表头Bold居中 / 表体居中

用法：
  python md_to_word_formal.py <input.md> [output.docx]
"""

import re
import sys
import os
from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ============================================================
# 格式参数常量（全部取自参考 docx 的精确值）
#   参考: 355600 EMU = 28pt, 406400 EMU = 32pt
#   中西文字体分离设置：eastAsia=中文字体, ascii/hAnsi=Times New Roman
# ============================================================

# --- 页面设置 ---
PAGE_WIDTH  = Cm(21.0)
PAGE_HEIGHT = Cm(29.7)
MARGIN_TOP    = Cm(3.70)
MARGIN_BOTTOM = Cm(3.50)
MARGIN_LEFT   = Cm(2.80)
MARGIN_RIGHT  = Cm(2.60)

# 字体常量（中西文分离）
# 注意：字体名必须与系统实际安装的字体名完全一致。
# 系统已安装：宋体(SimSun)、黑体(SimHei)、楷体(KaiTi)、仿宋(FangSong)、Times New Roman。
# 参考文档使用 方正X体 系列但该系统未安装，Word 打开参考文档时通过字体替换表渲染；
# python-docx 写入不存在的字体名会导致 Word 回退到 MS Gothic（日语默认字体）。
FONT_WESTERN = 'Times New Roman'     # 西文/数字统一使用
FONT_TITLE   = '宋体'               # 文档标题（中文：宋体，替代方正小标宋简体）
FONT_HEI     = '黑体'               # Heading 1（中文：黑体，替代方正黑体简体）
FONT_KAITI   = '楷体'               # Heading 2（中文）
FONT_SONG     = '宋体'               # 正文、Heading 3、表格（中文）
FONT_FANG    = '仿宋'               # Heading 4（中文：仿宋，替代方正仿宋简体）

# --- 文档标题（#）---
TITLE_EA_FONT     = FONT_TITLE
TITLE_SIZE        = Pt(20)
TITLE_BOLD        = True
TITLE_ALIGN       = WD_ALIGN_PARAGRAPH.CENTER
TITLE_LINE_SPACE  = Pt(32)          # EXACTLY 32pt (406400 EMU)
TITLE_SPACE_AFTER = Pt(18.05)

# --- Heading 1（## 一、二、…）---
H1_EA_FONT      = FONT_HEI
H1_SIZE         = Pt(16)
H1_BOLD         = True
H1_ALIGN        = WD_ALIGN_PARAGRAPH.LEFT
H1_LINE_SPACE   = Pt(28)           # EXACTLY 28pt (355600 EMU)
H1_SPACE_BEFORE = Pt(0)
H1_SPACE_AFTER  = Pt(0)

# --- Heading 2（###（一）（二）…）---
H2_EA_FONT      = FONT_KAITI
H2_SIZE         = Pt(16)
H2_BOLD         = True
H2_ALIGN        = WD_ALIGN_PARAGRAPH.LEFT
H2_LINE_SPACE   = Pt(28)
H2_SPACE_BEFORE = Pt(0)
H2_SPACE_AFTER  = Pt(0)

# --- Heading 3（#### 1. 2. …）---
H3_EA_FONT       = FONT_SONG
H3_SIZE          = Pt(14)
H3_BOLD          = True
H3_ALIGN         = WD_ALIGN_PARAGRAPH.LEFT
H3_LINE_SPACE    = Pt(28)
H3_SPACE_BEFORE  = Pt(0)
H3_SPACE_AFTER   = Pt(0)
H3_FIRST_INDENT  = Cm(1.42)

# --- Heading 4（#####（1）（2）…）---
H4_EA_FONT       = FONT_FANG
H4_SIZE          = Pt(14)
H4_BOLD          = True
H4_ALIGN         = WD_ALIGN_PARAGRAPH.LEFT
H4_LINE_SPACE    = Pt(28)
H4_SPACE_BEFORE  = Pt(9)
H4_SPACE_AFTER   = Pt(0)
H4_FIRST_INDENT  = Cm(1.42)

# --- 正文 ---
BODY_EA_FONT      = FONT_SONG
BODY_SIZE         = Pt(14)
BODY_ALIGN        = WD_ALIGN_PARAGRAPH.JUSTIFY
BODY_LINE_SPACE   = Pt(28)           # EXACTLY 28pt (355600 EMU)
BODY_FIRST_INDENT = Cm(1.13)

# --- 表格 ---
TBL_EA_FONT      = FONT_SONG
TBL_HEADER_SIZE  = Pt(12)
TBL_HEADER_BOLD  = True
TBL_HEADER_ALIGN = WD_ALIGN_PARAGRAPH.CENTER
TBL_BODY_SIZE    = Pt(12)
TBL_BODY_ALIGN   = WD_ALIGN_PARAGRAPH.CENTER


# ============================================================
# 编号辅助函数
# ============================================================

def _chinese_numeral(n):
    """整数转中文数字：1→一, 2→二, ..., 12→十二, 99→九十九。"""
    digits = ['', '一', '二', '三', '四', '五', '六', '七', '八', '九']
    if n < 1:
        return ''
    if n <= 9:
        return digits[n]
    if n <= 19:
        return '十' + (digits[n - 10] if n > 10 else '')
    tens = n // 10
    ones = n % 10
    return digits[tens] + '十' + (digits[ones] if ones > 0 else '')


def _circled_number(n):
    """整数转圈号：1→①, 2→②, ..., 暂支持 1-20，超出回退 (21) 等。"""
    circled = '①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑰⑱⑲⑳'
    if 1 <= n <= len(circled):
        return circled[n - 1]
    return f'({n})'


def _format_list_prefix(level, index):
    """根据标题层级生成列表编号前缀。

    #    (文档标题)  → 一、二、三、…
    ##   (Heading 1) → （一）（二）（三）…
    ###  (Heading 2) → 1. 2. 3. …
    #### (Heading 3) → （1）（2）（3）…
    #####(Heading 4) → ① ② ③ …
    """
    if level == 5:
        return _circled_number(index)
    elif level == 4:
        return f'（{index}）'
    elif level == 3:
        return f'{index}. '
    elif level == 2:
        return f'（{_chinese_numeral(index)}）'
    else:  # level == 1 or default
        return f'{_chinese_numeral(index)}、'


# ============================================================
# 辅助函数
# ============================================================

def _set_run_font(run, ea_font, font_size, bold=False):
    """设置单个 run 的中西文字体（eastAsia=中文字体, ascii/hAnsi=Times New Roman）。

    同时清除 rFonts 中可能存在的主题属性残留，仅保留显式字体名。"""
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
        rPr.insert(0, rFonts)
    else:
        # 清除可能已有的主题属性，防止与显式字体名冲突
        for theme_attr in ('w:asciiTheme', 'w:eastAsiaTheme', 'w:hAnsiTheme', 'w:cstheme'):
            attr_qn = qn(theme_attr)
            if attr_qn in rFonts.attrib:
                del rFonts.attrib[attr_qn]

    rFonts.set(qn('w:ascii'), FONT_WESTERN)
    rFonts.set(qn('w:hAnsi'), FONT_WESTERN)
    rFonts.set(qn('w:eastAsia'), ea_font)
    run.font.name = FONT_WESTERN
    run.font.size = font_size
    run.font.bold = bold


def _set_paragraph_fmt(paragraph, ea_font, font_size, bold, alignment,
                        line_spacing, space_before=None, space_after=None,
                        first_line_indent=None):
    """统一设置段落格式和全部 run 的中西文字体。"""
    pf = paragraph.paragraph_format
    pf.alignment = alignment
    if line_spacing is not None:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = line_spacing
    if space_before is not None:
        pf.space_before = space_before
    if space_after is not None:
        pf.space_after = space_after
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent

    for run in paragraph.runs:
        _set_run_font(run, ea_font, font_size, bold)


def _apply_inline_format(run, bold=False, italic=False):
    """对单个 run 设置内联格式（加粗/斜体），保留已有字体设置。"""
    if bold:
        run.font.bold = True
    if italic:
        run.font.italic = True


def _configure_heading_style(style, ea_font, font_size, bold, color_rgb,
                              alignment, line_spacing, space_before, space_after,
                              first_line_indent=None):
    """配置 Word 标题样式的中西文字体与段落格式。

    关键：python-docx 默认模板的 Heading 样式 rFonts 中带 asciiTheme/eastAsiaTheme/
    hAnsiTheme 主题属性。这些主题属性与显式字体名共存时会导致 Word 字体解析异常，
    必须清除主题属性，仅保留显式 ascii/hAnsi/eastAsia。
    """
    _style_rPr = style.element.get_or_add_rPr()
    _style_rFonts = _style_rPr.find(qn('w:rFonts'))
    if _style_rFonts is None:
        _style_rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
        _style_rPr.insert(0, _style_rFonts)

    # 清除主题字体属性（它们与显式字体名共存会导致 Word 回退到 MS Gothic）
    for theme_attr in ('w:asciiTheme', 'w:eastAsiaTheme', 'w:hAnsiTheme', 'w:cstheme'):
        attr_qn = qn(theme_attr)
        if attr_qn in _style_rFonts.attrib:
            del _style_rFonts.attrib[attr_qn]

    _style_rFonts.set(qn('w:ascii'), FONT_WESTERN)
    _style_rFonts.set(qn('w:hAnsi'), FONT_WESTERN)
    _style_rFonts.set(qn('w:eastAsia'), ea_font)
    style.font.name = FONT_WESTERN
    style.font.size = font_size
    style.font.bold = bold
    style.font.italic = False
    if color_rgb:
        style.font.color.rgb = color_rgb

    pf = style.paragraph_format
    pf.alignment = alignment
    if line_spacing is not None:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = line_spacing
    pf.space_before = space_before
    pf.space_after = space_after
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent


def _setup_document_styles(doc):
    """在文档对象上一次性配置所有标题样式（中英文字体分离）。"""
    style = doc.styles['Heading 1']
    _configure_heading_style(style, H1_EA_FONT, H1_SIZE, H1_BOLD, RGBColor(0, 0, 0),
                              H1_ALIGN, H1_LINE_SPACE, H1_SPACE_BEFORE, H1_SPACE_AFTER)

    style = doc.styles['Heading 2']
    _configure_heading_style(style, H2_EA_FONT, H2_SIZE, H2_BOLD, RGBColor(0, 0, 0),
                              H2_ALIGN, H2_LINE_SPACE, H2_SPACE_BEFORE, H2_SPACE_AFTER)

    style = doc.styles['Heading 3']
    _configure_heading_style(style, H3_EA_FONT, H3_SIZE, H3_BOLD, RGBColor(0, 0, 0),
                              H3_ALIGN, H3_LINE_SPACE, H3_SPACE_BEFORE, H3_SPACE_AFTER,
                              first_line_indent=H3_FIRST_INDENT)

    style = doc.styles['Heading 4']
    _configure_heading_style(style, H4_EA_FONT, H4_SIZE, H4_BOLD, RGBColor(0, 0, 0),
                              H4_ALIGN, H4_LINE_SPACE, H4_SPACE_BEFORE, H4_SPACE_AFTER,
                              first_line_indent=H4_FIRST_INDENT)


def _set_document_default_fonts(doc):
    """设置 docDefaults 显式字体（清除主题字体引用）。

    这是防治 MS Gothic 的关键防线：python-docx 默认模板的 docDefaults
    仅含主题引用（minorEastAsia / minorHAnsi），无显式字体名。
    当样式链中任何环节字体解析失败时，Word 会沿链回退到 docDefaults，
    若 docDefaults 也是主题引用则最终回退到 MS Gothic。

    与参考文档对齐：ascii=Times New Roman, eastAsia=宋体, hAnsi=Times New Roman。
    """
    styles_elem = doc.styles.element

    # 获取或创建 docDefaults
    doc_defaults = styles_elem.find(qn('w:docDefaults'))
    if doc_defaults is None:
        doc_defaults = parse_xml(f'<w:docDefaults {nsdecls("w")} />')
        styles_elem.insert(0, doc_defaults)

    rPr_default = doc_defaults.find(qn('w:rPrDefault'))
    if rPr_default is None:
        rPr_default = parse_xml(f'<w:rPrDefault {nsdecls("w")} />')
        doc_defaults.append(rPr_default)

    rPr = rPr_default.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")} />')
        rPr_default.append(rPr)

    # 移除旧的 rFonts（含主题属性）
    old_rFonts = rPr.find(qn('w:rFonts'))
    if old_rFonts is not None:
        rPr.remove(old_rFonts)

    # 创建新的 rFonts: 仅显式字体，无主题属性
    new_rFonts = parse_xml(
        f'<w:rFonts {nsdecls("w")} '
        f'w:ascii="{FONT_WESTERN}" '
        f'w:hAnsi="{FONT_WESTERN}" '
        f'w:eastAsia="{FONT_SONG}" '
        f'w:cs="{FONT_WESTERN}" />'
    )
    rPr.insert(0, new_rFonts)


def add_heading_paragraph(doc, text, level):
    """根据标题层级创建段落，使用 Word 内置标题样式和中英文字体分离。"""
    if level == 1:
        p = doc.add_paragraph()
        run = p.add_run(text)
        _set_paragraph_fmt(p, TITLE_EA_FONT, TITLE_SIZE, TITLE_BOLD, TITLE_ALIGN,
                           TITLE_LINE_SPACE, space_after=TITLE_SPACE_AFTER)
        return p

    style_map = {2: 'Heading 1', 3: 'Heading 2', 4: 'Heading 3', 5: 'Heading 4'}
    style_name = style_map.get(level, 'Normal')
    p = doc.add_paragraph(text, style=style_name)
    return p


def add_body_paragraph_with_inline(doc, text):
    """添加正文段落，支持 **加粗** 和 *斜体* 内联格式。"""
    p = doc.add_paragraph()
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|[^*]+)')
    parts = pattern.findall(text)
    for part in parts:
        full, bold_text, italic_text = part
        if full.startswith('**') and full.endswith('**'):
            run = p.add_run(bold_text)
            run.font.bold = True
        elif full.startswith('*') and full.endswith('*') and not full.startswith('**'):
            run = p.add_run(italic_text)
            run.font.italic = True
        else:
            run = p.add_run(full)
    _set_paragraph_fmt(p, BODY_EA_FONT, BODY_SIZE, False, BODY_ALIGN,
                       BODY_LINE_SPACE, first_line_indent=BODY_FIRST_INDENT)
    return p


def add_list_item_paragraph(doc, text, level, index):
    """添加列表项段落，根据章节层级使用对应的公文编号格式。

    Args:
        level: 当前所在章节的标题层级（1-5），对应 # 到 #####
        index: 当前编号组内的序号（1-based）
    """
    p = doc.add_paragraph()
    prefix = _format_list_prefix(level, index)
    p.add_run(prefix)
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|[^*]+)')
    parts = pattern.findall(text)
    for part in parts:
        full, bold_text, italic_text = part
        if full.startswith('**') and full.endswith('**'):
            run = p.add_run(bold_text)
            run.font.bold = True
        elif full.startswith('*') and full.endswith('*') and not full.startswith('**'):
            run = p.add_run(italic_text)
            run.font.italic = True
        else:
            run = p.add_run(full)
    _set_paragraph_fmt(p, BODY_EA_FONT, BODY_SIZE, False, BODY_ALIGN,
                       BODY_LINE_SPACE, first_line_indent=BODY_FIRST_INDENT)
    return p


def add_table_from_md(doc, lines):
    """将 Markdown 表格转换为 Word 三线表（顶线/底线 1.5pt，栏目线 0.75pt，无左右竖线）。"""
    header_line = lines[0]
    data_lines = lines[2:]

    headers = [h.strip() for h in header_line.strip('|').split('|')]
    num_cols = len(headers)

    data = []
    for dl in data_lines:
        cells = [c.strip() for c in dl.strip('|').split('|')]
        while len(cells) < num_cols:
            cells.append('')
        data.append(cells[:num_cols])

    table = doc.add_table(rows=1 + len(data), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # --- 构建单元格内容 ---
    def _fill_cell(cell, text, bold):
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = TBL_BODY_ALIGN
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = p.add_run(text)
        _set_run_font(run, TBL_EA_FONT, TBL_HEADER_SIZE if bold else TBL_BODY_SIZE, bold)

    for ci, h in enumerate(headers):
        _fill_cell(table.rows[0].cells[ci], h, bold=True)

    for ri, row_data in enumerate(data):
        for ci, val in enumerate(row_data):
            _fill_cell(table.rows[ri + 1].cells[ci], val, bold=False)

    # --- 三线表边框 XML ---
    # 清除所有默认边框，再逐行设置
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")} />')

    # 移除可能存在的旧边框设置
    for old_borders in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old_borders)

    borders_xml = (
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top    w:val="single" w:sz="12" w:space="0" w:color="000000" />'
        f'  <w:left   w:val="none"  w:sz="0"  w:space="0" w:color="auto" />'
        f'  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000" />'
        f'  <w:right  w:val="none"  w:sz="0"  w:space="0" w:color="auto" />'
        f'  <w:insideH w:val="none" w:sz="0"  w:space="0" w:color="auto" />'
        f'  <w:insideV w:val="none" w:sz="0"  w:space="0" w:color="auto" />'
        f'</w:tblBorders>'
    )
    tblPr.append(parse_xml(borders_xml))

    # 表头行底部边框（栏目线 0.75pt → w:sz="6"）
    for cell in table.rows[0].cells:
        tcPr = cell._tc.get_or_add_tcPr()
        for old_brd in tcPr.findall(qn('w:tcBorders')):
            tcPr.remove(old_brd)
        tcBorders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000" />'
            f'</w:tcBorders>'
        )
        tcPr.append(tcBorders)

    return table


def parse_markdown(md_text):
    """
    解析 Markdown 文本为结构化的 token 列表。

    返回 list of tuples: (type, content, extra)
      type: 'heading', 'body', 'list_item', 'blank', 'table'
      content: 文本内容
      extra: {'level': 1-5} for heading; {'level': int, 'index': int} for list_item

    注意：
      - 连续空行合并为单个 'blank' token。
      - 原有的有序/无序列表统一转为公文多级编号，编号格式取决于所在章节标题层级。
      - 列表项之间无空行视为同一编号组连续编号；遇空行或非列表内容重置编号。
    """
    lines = md_text.split('\n')
    tokens = []

    current_heading_level = 1   # 默认等同于文档标题级
    list_counter = 0
    in_list = False

    i = 0
    while i < len(lines):
        line = lines[i]
        rstripped = line.rstrip()

        # 空行处理：合并连续空行
        if not rstripped.strip():
            blank_count = 1
            i += 1
            while i < len(lines) and not lines[i].rstrip().strip():
                blank_count += 1
                i += 1

            # Peek ahead: 检查下一个非空行是否仍是列表项
            next_is_list = False
            if i < len(lines):
                nxt = lines[i].rstrip()
                next_is_list = bool(
                    re.match(r'^[\s]*[-*+]\s+', nxt) or
                    re.match(r'^[\s]*(\d+)\.\s+', nxt)
                )

            if in_list and next_is_list:
                tokens.append(('blank', '', None))
            else:
                tokens.append(('blank', '', None))
                in_list = False
                list_counter = 0
            continue

        # 标题
        heading_match = re.match(r'^(#{1,5})\s+(.+)', rstripped)
        if heading_match:
            level = len(heading_match.group(1))
            content = heading_match.group(2).strip()
            tokens.append(('heading', content, {'level': level}))
            current_heading_level = level
            in_list = False
            list_counter = 0
            i += 1
            continue

        # Markdown 表格
        if rstripped.startswith('|') and rstripped.endswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].rstrip().startswith('|') and lines[i].rstrip().endswith('|'):
                table_lines.append(lines[i].rstrip())
                i += 1
            tokens.append(('table', table_lines, None))
            in_list = False
            list_counter = 0
            continue

        # 无序列表（- / * / +）
        ul_match = re.match(r'^[\s]*[-*+]\s+(.+)', rstripped)
        if ul_match:
            if not in_list:
                list_counter = 0
                in_list = True
            list_counter += 1
            content = ul_match.group(1)
            tokens.append(('list_item', content, {'level': current_heading_level, 'index': list_counter}))
            i += 1
            continue

        # 有序列表（1. / 2. 等）
        ol_match = re.match(r'^[\s]*(\d+)\.\s+(.+)', rstripped)
        if ol_match:
            if not in_list:
                list_counter = 0
                in_list = True
            list_counter += 1
            content = ol_match.group(2)
            tokens.append(('list_item', content, {'level': current_heading_level, 'index': list_counter}))
            i += 1
            continue

        # 普通段落（正文）
        tokens.append(('body', rstripped.strip(), None))
        in_list = False
        list_counter = 0
        i += 1

    return tokens


def convert_md_to_docx(md_text, output_path):
    """将 Markdown 文本转换为符合排版规范的 docx 文件。"""
    doc = Document()

    # --- 页面设置 ---
    section = doc.sections[0]
    section.page_width  = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.top_margin    = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin   = MARGIN_LEFT
    section.right_margin  = MARGIN_RIGHT

    # --- 设置文档默认字体（必须在配置标题样式之前）---
    _set_document_default_fonts(doc)

    # --- 配置内置标题样式 ---
    _setup_document_styles(doc)

    # --- 解析 ---
    tokens = parse_markdown(md_text)

    # --- 剔除文件开头和结尾的 blank token ---
    while tokens and tokens[0][0] == 'blank':
        tokens.pop(0)
    while tokens and tokens[-1][0] == 'blank':
        tokens.pop()

    # --- 遍历生成（连续空行已合并为单个 blank，直接跳过即可）---
    for ttype, content, extra in tokens:
        if ttype == 'heading':
            add_heading_paragraph(doc, content, extra['level'])
        elif ttype == 'body':
            add_body_paragraph_with_inline(doc, content)
        elif ttype == 'list_item':
            add_list_item_paragraph(doc, content, extra['level'], extra['index'])
        elif ttype == 'table':
            add_table_from_md(doc, content)
        # 'blank' token 直接跳过——段落间间距由行距和段前/段后设置处理

    doc.save(output_path)
    return doc


# ============================================================
# 命令行入口
# ============================================================
if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("用法: python md_to_word_formal.py <input.md> [output.docx]")
        sys.exit(1)

    md_path = sys.argv[1]
    if not os.path.exists(md_path):
        print(f"错误: 文件不存在 - {md_path}")
        sys.exit(1)

    if len(sys.argv) >= 3:
        out_path = sys.argv[2]
    else:
        base = os.path.splitext(md_path)[0]
        out_path = base + '.docx'

    with open(md_path, 'r', encoding='utf-8') as f:
        md_text = f.read()

    convert_md_to_docx(md_text, out_path)
    print(f"转换完成: {out_path}")
